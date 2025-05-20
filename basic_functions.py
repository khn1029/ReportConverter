from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml import OxmlElement
from docx.shared import Inches
import os
import zipfile
from PIL import Image
import io
import shutil
import pandas as pd
import numpy as np
import hashlib
from docx.oxml.ns import qn
import imagehash
from PIL import Image as PILImage
import matplotlib.pyplot as plt
import re
import copy
from pandas.plotting import table
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from openpyxl.utils.dataframe import dataframe_to_rows
from html2image import Html2Image
import tempfile
import streamlit as st
from pathlib import Path
import zipfile
import time
import lxml.etree as ET

########### EXTRACT CAPTIONS ELEMENTS #################
# Define namespaces
from lxml import etree

# WordprocessingML namespace map
W_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def get_first_df_row(df): # take the first row of a filtered df
    if len(df)>0:
        return df.iloc[[0],:]
    else:
        return None

def get_elements_after_phrase(doc, trigger_phrase='executive summary'):
    found = False
    doc_copy = copy.deepcopy(doc)
    body = doc_copy.element.body
    collected = []

    for el in body.iterchildren():
        texts = el.findall('.//w:t', namespaces=W_NS)
        combined = ''.join([t.text.lower() for t in texts if t.text])

        if not found and trigger_phrase.lower() in combined:
            found = True
            continue

        if found:
            el_copy = copy.deepcopy(el)
            collected.append(el_copy)

    return collected


def extract_caption_sentence(text):
    # Just return full line for now (or you can customize)
    return text.strip()


def get_caption_after(doc,phrase=''):
    elements = get_elements_after_phrase(doc,phrase)
    results = []
    i = 0

    while i < len(elements):
        el = elements[i]
        tag = el.tag

        is_table = tag.endswith('tbl')
        is_image = tag.endswith('p') and (el.xpath('.//w:drawing') or el.xpath('.//w:pict'))

        if is_table or is_image:
            caption = None
            if i + 1 < len(elements):
                next_el = elements[i + 1]
                if next_el.tag.endswith('p'):
                    text_elems = next_el.xpath('.//w:t')
                    text = ''.join([t.text for t in text_elems if t.text])
                    caption = extract_caption_sentence(text)

            results.append(caption)
            i += 2
        else:
            i += 1

    return results


def get_caption_before(doc,phrase=''):
    elements = get_elements_after_phrase(doc,phrase)  # Get only relevant part of the doc
    results = []
    i = 0

    while i < len(elements):
        el = elements[i]
        tag = el.tag

        is_table = tag.endswith('tbl')
        is_image = tag.endswith('p') and (
                el.xpath('.//w:drawing') or el.xpath('.//w:pict')
        )

        if is_table or is_image:
            caption = None

            # Look one element before for caption
            if i - 1 >= 0:
                prev_el = elements[i - 1]
                if prev_el.tag.endswith('p'):
                    text_elems = prev_el.findall('.//w:t', namespaces={
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    text = ''.join([t.text for t in text_elems if t.text])
                    caption = extract_caption_sentence(text)

            results.append(caption)
            i += 1
        else:
            i += 1

    return results


def get_caption_both_side(doc,phrase):  # get the sentence before and after the elements
    caption_before = get_caption_before(doc,phrase)
    caption_after = get_caption_after(doc,phrase)
    return caption_before, caption_after



############## Extract table ##############

def convert_to_float(cell):
    try:
        return float(cell)
    except ValueError:
        return cell


# find first index of numeric to define multilevel index
def remove_unique_col(df):
    to_keep_col = []
    for i in range(len(df.columns)):
        if df.iloc[:, i].nunique() > 1:
            to_keep_col.extend([i])
    df_non_dup = df.iloc[:, to_keep_col]
    return df_non_dup


# get the level for multilevel index
def get_n_level(df):
    n_max = 0  ### only keep the largest index number
    for i in range(len(df.columns)):
        col_name = df.columns[i]
        cell_val = df.iloc[0, i]  # just to make column value different
        n = 0
        while cell_val == col_name:  # part of multilevel
            n += 1
            cell_val = df.iloc[n, i]
        if n_max <= n:
            n_max = n
    return n_max


def make_multilevel_header(df, n_header_rows):
    """
    Converts the first `n_header_rows` of `df` into a MultiIndex column header.
    Returns the cleaned dataframe with multi-level columns.
    """
    if n_header_rows > 0:
        # Step 1: Extract the first `n` rows as additional header levels
        sub_levels = df.iloc[:n_header_rows].astype(str).values

        # Step 2: Original column names as top-level
        top_level = df.columns.tolist()

        # Step 3: Combine into MultiIndex tuples
        tuples = [
            (top_level[i], *[sub_levels[j][i] for j in range(n_header_rows)])
            for i in range(len(top_level))
        ]
        multi_index = pd.MultiIndex.from_tuples(tuples)

        # Step 4: Drop header rows from data
        df_cleaned = df.iloc[n_header_rows:].reset_index(drop=True)

        # Step 5: Assign new multi-level columns
        df_cleaned.columns = multi_index
        return df_cleaned
    else:
        return df.copy()


#### Extract list of tables and figures
def extract_toc_from_docx(doc):
    # Load document
    # Combine all paragraph text
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # Get everything after "List of Figures"

    # Extract lines matching the format: Figure X.X<TAB>title<TAB>page
    # Use regex to handle tabs or multiple spaces
    figure_lines = re.findall(r'(Figure\s+\d+\.\d+.*?)\t+(.*?)\t+(\d+)', full_text)

    # Format nicely
    figure_dict = [{
        'label': match[0],
        'title': match[1],
        'page': int(match[2])
    } for match in figure_lines]

    ### TABLE
    table_lines = re.findall(r'(Table\s+\d+\.\d+.*?)\t+(.*?)\t+(\d+)', full_text)

    # Format nicely
    table_dict = [{
        'label': match[0],
        'title': match[1],
        'page': int(match[2])
    } for match in table_lines]

    # concat
    figure_dict.extend(table_dict)
    caption_df = pd.DataFrame(figure_dict)
    return caption_df



######### EXTRACT BOTH TABLES AND IMAGES
def get_df_from_tables(doc_table):
    df = pd.DataFrame(doc_table[1:], columns=doc_table[0]) if len(doc_table) > 1 else pd.DataFrame(doc_table)
    #### Need to check if true table
    #### Process the table
    df = df.map(convert_to_float)  # convert to float
    df = remove_unique_col(df)  # remove unique columns
    n_level = get_n_level(df)
    df = make_multilevel_header(df, n_level)
    return df


def extract_objects_from_doc(doc, phrase=""):
    # doc = Document(docx_path)
    object_list = list()
    # object_type

    # Step 1: Find the index where "Table of Contents" appears
    target_index = -1
    if phrase !='': # values is given
        for i, para in enumerate(doc.paragraphs):
            if phrase.upper() in para.text.upper():
                target_index = i
                break

    # Step 2: Start collecting tables after that point
    table_index = 0
    index = 0
    para_index = 0
    rels = doc.part._rels
    for block in doc.element.body.iterchildren():
        tag = block.tag

        # Track paragraph index
        if tag.endswith('p'):
            para_index += 1

            # Check for images in the paragraph
            if para_index > target_index and (block.xpath('.//w:drawing') or block.xpath('.//w:pict')):
                blips = block.xpath('.//a:blip')
                for blip in blips:
                    rId = blip.get(qn('r:embed'))
                    if rId and rId in rels:
                        image_part = rels[rId].target_part
                        image_bytes = image_part.blob
                        output = {'index': index, 'object': 'image', 'data': image_bytes}
                        object_list.append(output)
                        index += 1

        # Check for tables
        elif tag.endswith('tbl'):
            if para_index > target_index:
                table = doc.tables[table_index]
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
                df = get_df_from_tables(table_rows)
                output = {'index': index, 'object': 'table', 'data': df}
                object_list.append(output)
                index += 1

            table_index += 1

    return object_list


#### Join with captions to get Object name
def get_object_captions(object_index, toc_df):
    ### Join from top to bottom
    toc_df_filtered=toc_df.copy()
    toc_df_added=pd.DataFrame()
    for i in range(len(object_index)):
        if object_index['object'][i]=='image': # look for captions at bottom
            caption =object_index['caption_after'][i]
            new_row=get_first_df_row(toc_df_filtered[toc_df_filtered['title']==caption])
            ## If there is no match => table
            if new_row is None:
                caption = object_index['caption_before'][i]
                new_row = get_first_df_row(toc_df_filtered[toc_df_filtered['title']==caption])
        else: # must be table
            caption = object_index['caption_before'][i]
            new_row = get_first_df_row(toc_df_filtered[toc_df_filtered['title']==caption])
        toc_df_added=pd.concat([toc_df_added,new_row])
        # update toc_df_filtered to prevent duplication
        toc_df_filtered=toc_df_filtered[toc_df_filtered['label']!=new_row['label'].values[0]]
    combined=pd.concat([object_index,toc_df_added],axis=1)

    return combined.sort_index()


#### Convert df to table
def get_images_from_df(df, browser_path=None):
    """
    Convert a (possibly multi-header) DataFrame to a PNG image and return it
    as a BytesIO object ready for Streamlit or Excel embedding.

    Parameters
    ----------
    df : pandas.DataFrame
    browser_path : str | Path | None
        Full path to a Chromium-based browser executable.
        • Leave None to auto-detect (works if Google Chrome / Chromium is on PATH).
        • Pass Edge path on Windows, e.g.
          r"C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe"
    """
    # 1️⃣ Style the DataFrame → HTML
    styled_html = (
        df.style.set_table_attributes(
            'style="border-collapse:collapse; background-color:white; '
            'color:black; font-family:sans-serif;"'
        )
        .set_table_styles(
            [{"selector": "th, td", "props": [("padding", "8px"), ("border", "1px solid black")]}]
        )
        .to_html()
    )

    # 2️⃣ Guaranteed temp folder
    tmp_dir = Path(tempfile.gettempdir()) / "df_snaps"
    tmp_dir.mkdir(exist_ok=True)

    # 3️⃣ html2image instance
    hti = Html2Image(output_path=str(tmp_dir))
    if browser_path:  # allow caller to specify Edge / Brave, etc.
        hti.browser_executable = str(browser_path)

    # 4️⃣ Render → PNG
    output_name = "df_rendered.png"
    saved_files = hti.screenshot(html_str=styled_html, save_as=output_name)
    
    if not saved_files:  # html2image returns [] on failure
        raise RuntimeError(
            "html2image could not render the DataFrame.\n"
            "• Verify a Chromium browser exists or pass its path via browser_path.\n"
            "• On Windows you can pass Edge, e.g. browser_path=r'C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe'"
        )
    image_path = Path(saved_files[0])  
    #image_path = tmp_dir / output_name

    # 5️⃣ Load into memory
    with Image.open(image_path).copy() as img:  # copy → closes file handle
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)

    return buf
    
def show_dataframe_as_table(df):
    # convert \n to next line
    # Replace \n with <br> in the DataFrame for display purposes
    df_display = df.replace('\n', '<br>', regex=True)
    # Convert to HTML (escape = False lets Streamlit render it as HTML)
    html_table = df_display.to_html(index=False, escape=False)
    st.markdown(html_table, unsafe_allow_html=True)


def export_report_to_excel(object_list, excel_path,progress_callback=None):
    """Exports a dictionary of DataFrames to Excel with each DataFrame as an image."""
    max_size = (800, 800)
    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        # Access the workbook from the Excel writer
        workbook = writer.book
        toc_sheet = workbook.create_sheet("List of Tables and Figures")
        toc_sheet['A1'] = 'List of Tables'
        toc_sheet['F1'] = 'List of Figures'
        tabl_index = 2
        fig_index = 2
        object_num=1
        for object_dict in object_list:
            # get all data
            df = object_dict['data']
            object_type = object_dict['object']
            label = object_dict['label']
            caption = object_dict['title']

            # Create a new sheet and add the image
            sheet = workbook.create_sheet(str(label))

            if object_type.startswith('tabl'):  # check if dataframe                           
                image = get_images_from_df(df)
                excel_image = ExcelImage(image)
                sheet.add_image(excel_image, 'B2')
            else:  # already image
                image = PILImage.open(io.BytesIO(df))
                image = image.resize(max_size, resample=PILImage.BILINEAR)
                # Save image to a temporary BytesIO object
                image_stream = io.BytesIO()
                image.save(image_stream, format='PNG')
                image_stream.seek(0)

                # Create an openpyxl Image object from the BytesIO stream
                img = ExcelImage(image_stream)
                sheet.add_image(img, 'B2')  # Add image starting at cell A1
            ### Streamlit notification
                # Notify progress
            if progress_callback:
                progress_callback(object_num, len(object_list))

            # Optional: Add a caption to the image
            sheet['A1'] = caption
            ## Update TOC
            if label.lower().startswith('tabl'):
                toc_sheet['A' + str(tabl_index)] = label
                toc_sheet['B' + str(tabl_index)] = caption  #update TOC
                toc_sheet['A' + str(tabl_index)].hyperlink = f"#'{sheet.title}'!A1"
                tabl_index += 1
            else:
                toc_sheet['F' + str(fig_index)] = label
                toc_sheet['G' + str(fig_index)] = caption
                toc_sheet['F' + str(fig_index)].hyperlink = f"#'{sheet.title}'!A1"
                fig_index += 1
            object_num+=1


def get_object_name(object_list, object_index):
    # Attach title & label to each object
    object_list_copy = object_list.copy()
    for idx in range(len(object_list)):
        label = object_index.loc[object_index['index'] == idx, 'label'].values[0]
        title = object_index.loc[object_index['index'] == idx, 'title'].values[0]
        object_list_copy[idx].update({'label': label, 'title': title})
    return object_list_copy


### FINAL LOAD
def load_object_list(doc,phrase='EXECUTIVE SUMMARY',placeholder=None):
    toc_df = extract_toc_from_docx(doc)
    if placeholder is not None:
        placeholder.write('Finished loading TOC')
        time.sleep(1)
    object_list = extract_objects_from_doc(doc,phrase)

    # Create summary dataframe
    object_index = pd.DataFrame([{k: v for k, v in j.items() if k != 'data'} for j in object_list])
    caption_before, caption_after = get_caption_both_side(doc,phrase)
    object_index['caption_before'] = caption_before
    object_index['caption_after'] = caption_after
    # add captions to object_list
    object_index = get_object_captions(object_index, toc_df)

    # Attach title & label to each object
    object_list = get_object_name(object_list, object_index)
    if placeholder is not None:
        placeholder.write('Finished loading all tables and figures')
        time.sleep(1)
    return object_list

